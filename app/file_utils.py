import pandas as pd
from docx import Document
import os
import re

def sanitizar_nombre_archivo(nombre):
    """Limpia el nombre del archivo de caracteres inválidos."""
    # Quitamos caracteres no válidos para nombres de archivos en Windows/Linux
    nombre = re.sub(r'[\\/*?:"<>|]', "", nombre)
    # Reemplazamos espacios y saltos de línea por guiones bajos
    nombre = re.sub(r'\s+', "_", nombre).strip("_")
    return nombre

def actualizar_nota_excel(ruta_excel, alumno_nombre, materia, nota):
    """Busca al alumno y actualiza su nota en el archivo Excel real."""
    if not os.path.exists(ruta_excel):
        return f"⚠️ Error: No encuentro el archivo de notas en {ruta_excel}"
    
    try:
        df = pd.read_excel(ruta_excel)
        # Aseguramos que la columna Alumno existe
        if 'Alumno' not in df.columns:
            return f"⚠️ Error: El Excel no tiene una columna llamada 'Alumno'. Las columnas son: {df.columns.tolist()}"
        
        # Buscamos una coincidencia parcial (por si solo dice el nombre o el apellido)
        mask = df['Alumno'].str.contains(alumno_nombre, case=False, na=False)
        
        if mask.any():
            # Actualizamos la nota en la primera coincidencia encontrada
            idx = df[mask].index[0]
            df.at[idx, 'Nota'] = nota
            df.at[idx, 'Materia'] = materia.capitalize()
            df.to_excel(ruta_excel, index=False)
            return f"✅ ¡Listo Profe! Le puse un {nota} a {df.at[idx, 'Alumno']} en {materia}."
        else:
            return f"❓ Profe, no encontré a nadie llamado '{alumno_nombre}' en la lista. ¿Puede verificar cómo está escrito?"
    except Exception as e:
        return f"❌ Hubo un problema con el Excel: {str(e)}"

def generar_word_planeamiento(datos, carpeta_salida="data/output"):
    """Crea un documento Word extenso y profesional."""
    tema_original = datos.get("tema", "Tema General")
    tema = sanitizar_nombre_archivo(tema_original)
    grado = datos.get("grado", "Grado no especificado")
    
    try:
        doc = Document()
        doc.add_heading('PLANIFICACIÓN DIDÁCTICA DE CLASE', 0)
        
        table = doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        table.cell(0, 0).text = f"Tema: {tema_original}"
        table.cell(0, 1).text = f"Grado: {grado}"
        table.cell(1, 0).text = "Institución: ____________________"
        table.cell(1, 1).text = "Fecha: ____/____/202X"

        doc.add_heading('I. OBJETIVOS DE APRENDIZAJE', level=1)
        objetivos = datos.get("objetivos", [])
        if isinstance(objetivos, list):
            for obj in objetivos:
                doc.add_paragraph(obj, style='List Bullet')
        else:
            doc.add_paragraph(str(objetivos))

        doc.add_heading('II. RECURSOS Y MATERIALES', level=1)
        recursos = datos.get("recursos", [])
        if isinstance(recursos, list):
            doc.add_paragraph(", ".join(recursos))
        else:
            doc.add_paragraph(str(recursos))

        doc.add_heading('III. SECUENCIA DIDÁCTICA', level=1)
        doc.add_heading('Inicio (Motivación)', level=2)
        doc.add_paragraph(datos.get("inicio", "Sin detalles especificados."))
        
        doc.add_heading('Desarrollo (Construcción)', level=2)
        doc.add_paragraph(datos.get("desarrollo", "Sin detalles especificados."))
        
        doc.add_heading('Cierre (Evaluación)', level=2)
        doc.add_paragraph(datos.get("cierre", "Sin detalles especificados."))

        doc.add_heading('IV. INDICADORES DE EVALUACIÓN', level=1)
        indicadores = datos.get("indicadores", [])
        if isinstance(indicadores, list):
            for ind in indicadores:
                doc.add_paragraph(ind, style='List Bullet')
        else:
            doc.add_paragraph(str(indicadores))

        if not os.path.exists(carpeta_salida):
            os.makedirs(carpeta_salida)
            
        archivo_nombre = f"Plan_Detallado_{tema}.docx"
        ruta_guardado = os.path.join(carpeta_salida, archivo_nombre)
        
        # Guardar el documento
        doc.save(ruta_guardado)
        return ruta_guardado, archivo_nombre
    except Exception as e:
        # En caso de error, intentamos un nombre genérico
        archivo_nombre = f"Plan_Emergencia_{hash(tema)}.docx"
        ruta_guardado = os.path.join(carpeta_salida, archivo_nombre)
        doc.save(ruta_guardado)
        return ruta_guardado, archivo_nombre
